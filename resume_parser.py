"""
Resume Parser Service for Job2Hire
Extracts structured candidate data from uploaded resumes using document parsing and AI
"""

import os
import re
import json
import logging
from datetime import datetime
from typing import Dict, List, Optional, Tuple
import pdfplumber
import PyPDF2
from docx import Document
from openai import OpenAI

# Initialize OpenAI client
openai_client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))

class ResumeParser:
    """
    Automated resume parsing service that extracts structured candidate data
    from PDF, DOCX, and other document formats
    """
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt']
        self.logger = logging.getLogger(__name__)
    
    def parse_resume(self, file_path: str, filename: str) -> Dict:
        """
        Main parsing function that extracts all candidate information from resume
        
        Args:
            file_path: Path to the uploaded resume file
            filename: Original filename for format detection
            
        Returns:
            Dict containing extracted candidate information
        """
        try:
            # Extract raw text from document
            raw_text = self._extract_text_from_document(file_path, filename)
            
            if not raw_text.strip():
                return {"error": "Could not extract text from document"}
            
            # Use AI to extract structured information
            structured_data = self._extract_structured_data(raw_text)
            
            # Clean and validate extracted data
            cleaned_data = self._clean_and_validate_data(structured_data)
            
            return {
                "success": True,
                "data": cleaned_data,
                "raw_text": raw_text[:1000] + "..." if len(raw_text) > 1000 else raw_text
            }
            
        except Exception as e:
            self.logger.error(f"Resume parsing failed: {e}")
            return {"error": f"Resume parsing failed: {str(e)}"}
    
    def _extract_text_from_document(self, file_path: str, filename: str) -> str:
        """Extract raw text from various document formats"""
        file_extension = os.path.splitext(filename)[1].lower()
        
        if file_extension == '.pdf':
            return self._extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return self._extract_text_from_docx(file_path)
        elif file_extension == '.doc':
            return self._extract_text_from_doc(file_path)
        elif file_extension == '.txt':
            return self._extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF using multiple methods for better accuracy"""
        text = ""
        
        # Method 1: Try pdfplumber first (better formatting)
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            self.logger.warning(f"pdfplumber failed: {e}")
        
        # Method 2: Fallback to PyPDF2 if pdfplumber fails
        if not text.strip():
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
            except Exception as e:
                self.logger.warning(f"PyPDF2 failed: {e}")
        
        return text.strip()
    
    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX files"""
        try:
            doc = Document(file_path)
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            return '\n'.join(text)
        except Exception as e:
            self.logger.error(f"DOCX extraction failed: {e}")
            return ""
    
    def _extract_text_from_doc(self, file_path: str) -> str:
        """Extract text from DOC files using fallback methods"""
        try:
            # First try reading as DOCX (sometimes .doc files are actually .docx)
            try:
                doc = Document(file_path)
                text = []
                for paragraph in doc.paragraphs:
                    text.append(paragraph.text)
                result = '\n'.join(text)
                if result.strip():
                    return result
            except:
                pass
            
            # Fallback: Try reading as plain text
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                    content = file.read()
                    # Basic cleaning for binary content
                    import re
                    # Remove non-printable characters but keep basic formatting
                    cleaned = re.sub(r'[^\x20-\x7E\n\r\t]', ' ', content)
                    # Remove excessive whitespace
                    cleaned = re.sub(r'\s+', ' ', cleaned)
                    if len(cleaned.strip()) > 50:  # Reasonable content length
                        return cleaned
            except:
                pass
            
            # Final fallback: Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1', errors='ignore') as file:
                    content = file.read()
                    import re
                    cleaned = re.sub(r'[^\x20-\x7E\n\r\t]', ' ', content)
                    cleaned = re.sub(r'\s+', ' ', cleaned)
                    if len(cleaned.strip()) > 50:
                        return cleaned
            except:
                pass
                
            return ""
            
        except Exception as e:
            self.logger.error(f"DOC extraction failed: {e}")
            return ""
    
    def _extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try different encodings
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                self.logger.error(f"TXT extraction failed: {e}")
                return ""
    
    def _extract_structured_data(self, raw_text: str) -> Dict:
        """Use OpenAI to extract structured information from raw text"""
        try:
            prompt = f"""
            Extract the following information from this resume text and return it as a JSON object:
            
            {{
                "personal_info": {{
                    "first_name": "string",
                    "last_name": "string", 
                    "email": "string",
                    "phone": "string",
                    "location": "string",
                    "linkedin_url": "string",
                    "portfolio_url": "string"
                }},
                "professional_summary": {{
                    "current_job_title": "string",
                    "bio": "string (2-3 sentences)",
                    "experience_years": "integer"
                }},
                "skills": ["skill1", "skill2", "skill3"],
                "work_experience": [
                    {{
                        "title": "string",
                        "company": "string",
                        "duration": "string",
                        "description": "string",
                        "start_date": "YYYY-MM or YYYY",
                        "end_date": "YYYY-MM or YYYY or null if current",
                        "current": "boolean"
                    }}
                ],
                "education": [
                    {{
                        "degree": "string (e.g., B.Tech, 12th, 10th, Masters, etc.)",
                        "institution": "string (school/college/university name)",
                        "year": "string (year of passing/completion)",
                        "field": "string (field of study/specialization)",
                        "grade": "string (percentage, CGPA, or grade if available)",
                        "board": "string (board/university name if different from institution)"
                    }}
                ],
                "certifications": [
                    {{
                        "name": "string",
                        "issuer": "string",
                        "year": "string"
                    }}
                ]
            }}
            
            Special Instructions for Education:
            - Look for education tables with columns like QUALIFICATION, BOARD/UNIVERSITY, SCHOOL/COLLEGE, YEAR, GRADE/PERCENTAGE
            - Extract each educational qualification as a separate entry (B.Tech, 12th, 10th, etc.)
            - Include both formal degrees and school qualifications
            - Capture grades, percentages, CGPA when available
            - If field of study is mentioned (like "Dairy Technology"), include it in the field
            - If board is different from institution (like CBSE, SHIATS), capture both
            
            Resume text:
            {raw_text}
            
            Important: Return only valid JSON. Pay special attention to tabular education data. If information is not found, use null for strings and empty arrays for lists.
            """
            
            response = openai_client.chat.completions.create(
                model="gpt-4o",  # newest OpenAI model
                messages=[
                    {"role": "system", "content": "You are an expert resume parser. Extract information accurately and return only valid JSON."},
                    {"role": "user", "content": prompt}
                ],
                response_format={"type": "json_object"},
                temperature=0.1
            )
            
            content = response.choices[0].message.content
            if content:
                return json.loads(content)
            else:
                raise Exception("Empty response from AI")
            
        except Exception as e:
            self.logger.error(f"AI extraction failed: {e}")
            # Fallback to regex-based extraction
            return self._fallback_extraction(raw_text)
    
    def _fallback_extraction(self, raw_text: str) -> Dict:
        """Fallback extraction using regex patterns when AI fails"""
        data = {
            "personal_info": {},
            "professional_summary": {},
            "skills": [],
            "work_experience": [],
            "education": [],
            "certifications": []
        }
        
        # Extract email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, raw_text)
        if emails:
            data["personal_info"]["email"] = emails[0]
        
        # Extract phone number
        phone_pattern = r'(\+?1[-.\s]?)?(\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4})'
        phones = re.findall(phone_pattern, raw_text)
        if phones:
            data["personal_info"]["phone"] = ''.join(phones[0])
        
        # Extract LinkedIn URL
        linkedin_pattern = r'(?:https?://)?(?:www\.)?linkedin\.com/in/[A-Za-z0-9-_.]+'
        linkedin_urls = re.findall(linkedin_pattern, raw_text)
        if linkedin_urls:
            data["personal_info"]["linkedin_url"] = linkedin_urls[0]
        
        # Extract skills (common programming languages and technologies)
        skills_pattern = r'\b(?:Python|JavaScript|Java|C\+\+|React|Node\.js|SQL|HTML|CSS|AWS|Docker|Git)\b'
        skills = list(set(re.findall(skills_pattern, raw_text, re.IGNORECASE)))
        data["skills"] = skills
        
        return data
    
    def _clean_and_validate_data(self, data: Dict) -> Dict:
        """Clean and validate extracted data"""
        cleaned_data = {}
        
        # Clean personal info
        personal_info = data.get("personal_info", {})
        cleaned_data["personal_info"] = {
            "first_name": self._clean_text(personal_info.get("first_name")),
            "last_name": self._clean_text(personal_info.get("last_name")),
            "email": self._validate_email(personal_info.get("email")),
            "phone": self._clean_phone(personal_info.get("phone")),
            "location": self._clean_text(personal_info.get("location")),
            "linkedin_url": self._validate_url(personal_info.get("linkedin_url")),
            "portfolio_url": self._validate_url(personal_info.get("portfolio_url"))
        }
        
        # Clean professional summary
        prof_summary = data.get("professional_summary", {})
        cleaned_data["professional_summary"] = {
            "current_job_title": self._clean_text(prof_summary.get("current_job_title")),
            "bio": self._clean_text(prof_summary.get("bio")),
            "experience_years": self._validate_integer(prof_summary.get("experience_years"))
        }
        
        # Clean skills
        skills = data.get("skills", [])
        cleaned_data["skills"] = [self._clean_text(skill) for skill in skills if skill]
        
        # Clean work experience
        work_exp = data.get("work_experience", [])
        cleaned_data["work_experience"] = [
            {
                "title": self._clean_text(exp.get("title")),
                "company": self._clean_text(exp.get("company")),
                "duration": self._clean_text(exp.get("duration")),
                "description": self._clean_text(exp.get("description")),
                "start_date": exp.get("start_date"),
                "end_date": exp.get("end_date"),
                "current": bool(exp.get("current", False))
            }
            for exp in work_exp if exp.get("title") and exp.get("company")
        ]
        
        # Clean education
        education = data.get("education", [])
        cleaned_data["education"] = [
            {
                "degree": self._clean_text(edu.get("degree")),
                "institution": self._clean_text(edu.get("institution")),
                "year": self._clean_text(edu.get("year")),
                "field": self._clean_text(edu.get("field"))
            }
            for edu in education if edu.get("degree") and edu.get("institution")
        ]
        
        # Clean certifications
        certifications = data.get("certifications", [])
        cleaned_data["certifications"] = [
            {
                "name": self._clean_text(cert.get("name")),
                "issuer": self._clean_text(cert.get("issuer")),
                "year": self._clean_text(cert.get("year"))
            }
            for cert in certifications if cert.get("name")
        ]
        
        return cleaned_data
    
    def _clean_text(self, text: str) -> Optional[str]:
        """Clean and normalize text"""
        if not text or str(text).lower() in ['null', 'none', 'n/a']:
            return None
        return str(text).strip()
    
    def _validate_email(self, email: str) -> Optional[str]:
        """Validate email format"""
        if not email:
            return None
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        if re.match(email_pattern, email):
            return email.lower()
        return None
    
    def _clean_phone(self, phone: str) -> Optional[str]:
        """Clean and format phone number"""
        if not phone:
            return None
        # Remove all non-digit characters
        digits = re.sub(r'\D', '', phone)
        if len(digits) == 10:
            return f"+1{digits}"
        elif len(digits) == 11 and digits.startswith('1'):
            return f"+{digits}"
        return phone
    
    def _validate_url(self, url: str) -> Optional[str]:
        """Validate URL format"""
        if not url:
            return None
        if not url.startswith(('http://', 'https://')):
            url = 'https://' + url
        return url
    
    def _validate_integer(self, value) -> Optional[int]:
        """Validate integer value"""
        if isinstance(value, int):
            return value
        if isinstance(value, str) and value.isdigit():
            return int(value)
        return None

# Utility functions for integration with Job2Hire
def parse_resume_file(file_path: str, filename: str) -> Dict:
    """
    Parse resume file and return structured data
    
    Args:
        file_path: Path to uploaded resume file
        filename: Original filename
        
    Returns:
        Dict containing parsed resume data or error message
    """
    parser = ResumeParser()
    return parser.parse_resume(file_path, filename)

def populate_user_profile(user, resume_data: Dict) -> bool:
    """
    Populate user profile with parsed resume data
    
    Args:
        user: User model instance
        resume_data: Parsed resume data dictionary
        
    Returns:
        bool: True if successful, False otherwise
    """
    try:
        if not resume_data.get("success") or not resume_data.get("data"):
            return False
        
        data = resume_data["data"]
        
        # Populate personal information
        personal_info = data.get("personal_info", {})
        if personal_info.get("first_name"):
            user.first_name = personal_info["first_name"]
        if personal_info.get("last_name"):
            user.last_name = personal_info["last_name"]
        if personal_info.get("email"):
            user.email = personal_info["email"]
        if personal_info.get("phone"):
            user.phone = personal_info["phone"]
        if personal_info.get("location"):
            user.location = personal_info["location"]
        if personal_info.get("linkedin_url"):
            user.linkedin_url = personal_info["linkedin_url"]
        if personal_info.get("portfolio_url"):
            user.portfolio_url = personal_info["portfolio_url"]
        
        # Populate professional summary
        prof_summary = data.get("professional_summary", {})
        if prof_summary.get("current_job_title"):
            user.job_title = prof_summary["current_job_title"]
        if prof_summary.get("bio"):
            user.bio = prof_summary["bio"]
        if prof_summary.get("experience_years"):
            user.experience_years = prof_summary["experience_years"]
        
        # Populate skills
        skills = data.get("skills", [])
        if skills:
            user.skills = json.dumps(skills)
        
        # Populate work experience
        work_experience = data.get("work_experience", [])
        if work_experience:
            user.experience = json.dumps(work_experience)
        
        # Populate education
        education = data.get("education", [])
        if education:
            user.education = json.dumps(education)
        
        # Populate certifications
        certifications = data.get("certifications", [])
        if certifications:
            user.certifications = json.dumps(certifications)
        
        return True
        
    except Exception as e:
        logging.error(f"Error populating user profile: {e}")
        return False